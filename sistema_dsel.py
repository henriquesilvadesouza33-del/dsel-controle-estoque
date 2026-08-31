import base64
import streamlit as st
import sqlite3
import pandas as pd
from datetime import datetime, date
import os
import io
import unicodedata
from docx import Document
from fpdf import FPDF
try:
    from PIL import Image
    PIL_DISPONIVEL = True
except ImportError:
    PIL_DISPONIVEL = False

# --- 1. CONFIGURAÇÕES E ESTILO ---
st.set_page_config(page_title="DSEL - GESTÃO NOVACAP", layout="wide")
DB_NAME = 'estoque_dsel_v3.db'
NOME_LOGO = "logo.novacap.png"
NOME_LOGO_PDF = "NOVACAP_BRASILIA_LOGO.png"

st.markdown("""
    <style>
    .stMetric { background-color: #1e2630; padding: 15px; border-radius: 10px; border-left: 5px solid #0056b3; }
    .stButton>button { width: 100%; font-weight: bold; border-radius: 8px; height: 3em; }
    .stDataFrame { background-color: #1e2630; border-radius: 10px; }
    .footer-henrique { 
        position: fixed; 
        bottom: 0; 
        left: 0; 
        width: 100%; 
        background-color: #0e1117; 
        color: #888; 
        text-align: center; 
        padding: 10px; 
        border-top: 1px solid #333; 
        font-size: 13px; 
        z-index: 999;
    }
    </style>
""", unsafe_allow_html=True)


def rodape_henrique():
    st.markdown("""
        <div class="footer-henrique">
            Este sistema foi desenvolvido por Henrique Silva de Souza
        </div>
    """, unsafe_allow_html=True)


# --- 2. BANCO DE DADOS ---
def conectar(): 
    return sqlite3.connect(DB_NAME)

def inicializar_banco():
    with conectar() as conn:
        c = conn.cursor()
        c.execute('''CREATE TABLE IF NOT EXISTS usuarios (
                        id INTEGER PRIMARY KEY AUTOINCREMENT, 
                        nome TEXT UNIQUE, 
                        senha TEXT, 
                        status TEXT)''')
        c.execute('''CREATE TABLE IF NOT EXISTS estoque (
                        id INTEGER PRIMARY KEY AUTOINCREMENT, 
                        item TEXT, 
                        cor TEXT, 
                        tamanho TEXT, 
                        modelo TEXT, 
                        qtd INTEGER, 
                        estado TEXT, 
                        nf TEXT, 
                        pessoa TEXT, 
                        data_hora TEXT, 
                        almoxarife TEXT)''')
        c.execute("INSERT OR IGNORE INTO usuarios (nome, senha, status) VALUES (?, ?, ?)", ('ADMIN', '123456', 'Ativo'))
        c.execute("UPDATE estoque SET modelo = 'BRIM' WHERE modelo IN ('Calça BRIM', 'brim', 'Brim')")
        c.execute("UPDATE estoque SET modelo = 'Padrão' WHERE modelo IN ('Manga Curta', 'Manga Longa')")
        conn.commit()

inicializar_banco()


# --- 3. FUNÇÕES DE APOIO ---
def exibir_logo(largura=200):
    if os.path.exists(NOME_LOGO): 
        st.image(NOME_LOGO, width=largura)

def consultar_saldo(item, cor, tamanho, modelo, estado):
    with conectar() as conn:
        res = conn.execute("SELECT SUM(qtd) FROM estoque WHERE item=? AND cor=? AND tamanho=? AND modelo=? AND estado=?", 
                           (item, cor, tamanho, modelo, estado)).fetchone()
        return res[0] if res and res[0] is not None else 0


# --- 4. FUNÇÕES DE EXPORTAÇÃO ---
def gerar_csv(df):
    return df.to_csv(index=False).encode('utf-8-sig')

def gerar_excel(df):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Dados')
    return output.getvalue()

def gerar_word(df, titulo="Relatório DSEL NOVACAP"):
    doc = Document()
    doc.add_heading(titulo, 0)
    doc.add_paragraph(f"Relatório gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value) if value is not None else ""
    doc_io = io.BytesIO()
    doc.save(doc_io)
    return doc_io.getvalue()

def remover_acentos(texto):
    """Remove acentos para compatibilidade com FPDF (latin-1)."""
    if not isinstance(texto, str):
        texto = str(texto)
    # Mapeamento explicito para caracteres acentuados do portugues
    mapeamento = {
        'ç': 'c', 'Ç': 'C',
        'ã': 'a', 'Ã': 'A',
        'õ': 'o', 'Õ': 'O',
        'á': 'a', 'Á': 'A',
        'é': 'e', 'É': 'E',
        'í': 'i', 'Í': 'I',
        'ó': 'o', 'Ó': 'O',
        'ú': 'u', 'Ú': 'U',
        'â': 'a', 'Â': 'A',
        'ê': 'e', 'Ê': 'E',
        'ô': 'o', 'Ô': 'O',
        'à': 'a', 'À': 'A',
        'ñ': 'n', 'Ñ': 'N',
    }
    for char, subst in mapeamento.items():
        texto = texto.replace(char, subst)
    # Remove marcas de acento restantes (NFD)
    return "".join(c for c in unicodedata.normalize('NFD', texto) if unicodedata.category(c) != 'Mn')


def gerar_pdf(df, titulo="Relatorio DSEL NOVACAP"):
    pdf = FPDF(orientation='L', unit='mm', format='A4')
    pdf.add_page()
    pdf.set_font("Arial", 'B', 12)
    pdf.cell(0, 10, remover_acentos(titulo), ln=1, align='C')
    pdf.set_font("Arial", '', 8)
    pdf.cell(0, 6, f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}", ln=1, align='C')
    pdf.ln(4)

    # Larguras específicas por coluna para evitar sobreposição
    larguras_padrao = {
        'id': 12, 'ID': 12,
        'Data/Hora': 28, 'data_hora': 28,
        'Item': 22, 'item': 22,
        'Modelo': 26, 'modelo': 26,
        'Cor': 18, 'cor': 18,
        'Tam.': 14, 'tamanho': 14,
        'Qtd': 12, 'qtd': 12,
        'Qtd. Retirada': 16,
        'Estado': 18, 'estado': 18,
        'NF': 20, 'nf': 20,
        'Origem/Destino': 48, 'pessoa': 48,
        'Reeducando': 48,
        'Almoxarife': 30, 'almoxarife': 30,
        'Saldo': 14, 'Saldo Atual': 16
    }

    n_cols = len(df.columns)
    largura_total = 277  # A4 landscape ~297mm - 10mm margem cada lado

    col_widths = []
    for col in df.columns:
        col_str = str(col)
        if col_str in larguras_padrao:
            col_widths.append(larguras_padrao[col_str])
        else:
            encontrado = False
            for key, val in larguras_padrao.items():
                if key.lower() in col_str.lower():
                    col_widths.append(val)
                    encontrado = True
                    break
            if not encontrado:
                col_widths.append(max(15, largura_total // n_cols))

    soma = sum(col_widths)
    if soma > largura_total:
        fator = largura_total / soma
        col_widths = [w * fator for w in col_widths]

    pdf.set_font("Arial", 'B', 7)
    for i, col in enumerate(df.columns):
        nome_col = remover_acentos(str(col))
        max_chars = int(col_widths[i] / 1.8)
        pdf.cell(col_widths[i], 6, nome_col[:max_chars], border=1, align='C')
    pdf.ln()

    pdf.set_font("Arial", '', 7)
    for index, row in df.iterrows():
        for i, val in enumerate(row):
            texto_celula = remover_acentos(val) if val is not None else ""
            max_chars = int(col_widths[i] / 1.6)
            texto_truncado = texto_celula[:max_chars] if len(texto_celula) > max_chars else texto_celula
            pdf.cell(col_widths[i], 5, texto_truncado, border=1, align='C')
        pdf.ln()
    return pdf.output(dest='S').encode('latin-1', errors='ignore')


def botoes_exportacao(df, nome_base="relatorio", titulo_pdf="Relatorio DSEL NOVACAP"):
    col1, col2, col3, col4 = st.columns(4)
    data_str = datetime.now().strftime('%Y%m%d_%H%M')
    with col1:
        st.download_button("📊 Excel", gerar_excel(df), f"{nome_base}_{data_str}.xlsx", 
                          "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    with col2:
        st.download_button("📝 Word", gerar_word(df, titulo_pdf), f"{nome_base}_{data_str}.docx", 
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    with col3:
        st.download_button("📄 PDF", gerar_pdf(df, titulo_pdf), f"{nome_base}_{data_str}.pdf", 
                          "application/pdf")
    with col4:
        st.download_button("📋 CSV", gerar_csv(df), f"{nome_base}_{data_str}.csv", 
                          "text/csv")


# --- COMPROVANTE DE ENTREGA (PDF INDIVIDUAL) ---
def gerar_comprovante_entrega(nome_recebedor, df_itens, almoxarife_nome):
    pdf = FPDF(orientation='P', unit='mm', format='A4')
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Logo NOVACAP no canto superior esquerdo
    logo_pdf = NOME_LOGO_PDF if os.path.exists(NOME_LOGO_PDF) else NOME_LOGO
    if os.path.exists(logo_pdf) and PIL_DISPONIVEL:
        try:
            img = Image.open(logo_pdf)
            if img.mode in ('RGBA', 'LA', 'P'):
                img = img.convert('RGB')
            temp_logo = "_temp_logo_novacap.png"
            img.save(temp_logo, format='PNG')
            pdf.image(temp_logo, x=10, y=8, w=30)
            if os.path.exists(temp_logo):
                os.remove(temp_logo)
        except Exception:
            pass

    # Cabecalho centralizado
    pdf.set_y(12)
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 8, "NOVACAP - DSEL".encode('latin1', 'replace').decode('latin1'), ln=1, align='C')
    pdf.set_font("Arial", 'B', 10)
    pdf.cell(0, 6, "DIVISÃO DE SEGURANÇA, LIMPEZA E CONSERVAÇÃO".encode('latin1', 'replace').decode('latin1'), ln=1, align='C')
    pdf.set_font("Arial", '', 10)
    pdf.cell(0, 6, "COMPROVANTE DE ENTREGA DE MATERIAL".encode('latin1', 'replace').decode('latin1'), ln=1, align='C')
    pdf.ln(2)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(6)

    # Dados do recebedor
    pdf.set_font("Arial", 'B', 10)
    pdf.cell(60, 8, "Recebedor:".encode('latin1', 'replace').decode('latin1'), ln=0)
    pdf.set_font("Arial", '', 10)
    pdf.cell(0, 8, str(nome_recebedor).encode('latin1', 'replace').decode('latin1'), ln=1)

    pdf.set_font("Arial", 'B', 10)
    pdf.cell(60, 8, "Data da Entrega:".encode('latin1', 'replace').decode('latin1'), ln=0)
    pdf.set_font("Arial", '', 10)
    pdf.cell(0, 8, date.today().strftime("%d/%m/%Y"), ln=1)

    pdf.set_font("Arial", 'B', 10)
    pdf.cell(60, 8, "Responsável pelo Lançamento:".encode('latin1', 'replace').decode('latin1'), ln=0)
    pdf.set_font("Arial", '', 10)
    pdf.cell(0, 8, str(almoxarife_nome).encode('latin1', 'replace').decode('latin1'), ln=1)
    pdf.ln(5)

    # Tabela de itens
    pdf.set_font("Arial", 'B', 10)
    pdf.set_fill_color(220, 220, 220)
    pdf.cell(60, 8, "Item".encode('latin1', 'replace').decode('latin1'), border=1, align='C', fill=True)
    pdf.cell(35, 8, "Modelo".encode('latin1', 'replace').decode('latin1'), border=1, align='C', fill=True)
    pdf.cell(30, 8, "Cor".encode('latin1', 'replace').decode('latin1'), border=1, align='C', fill=True)
    pdf.cell(25, 8, "Tam.".encode('latin1', 'replace').decode('latin1'), border=1, align='C', fill=True)
    pdf.cell(20, 8, "Qtd".encode('latin1', 'replace').decode('latin1'), border=1, align='C', fill=True)
    pdf.cell(20, 8, "Estado".encode('latin1', 'replace').decode('latin1'), border=1, align='C', fill=True)
    pdf.ln()

    pdf.set_font("Arial", '', 9)
    for _, row in df_itens.iterrows():
        # Substitui textos sem acento vindos do banco para o padrão correto
        item_texto = str(row.get('Item', '')).replace('Calca', 'Calça')
        estado_texto = str(row.get('Estado', '')).replace('Regulamentar', 'Regulamentar') # Se houver algum no banco, mude aqui

        pdf.cell(60, 7, item_texto[:28].encode('latin1', 'replace').decode('latin1'), border=1, align='L')
        pdf.cell(35, 7, str(row.get('Modelo', ''))[:16].encode('latin1', 'replace').decode('latin1'), border=1, align='L')
        pdf.cell(30, 7, str(row.get('Cor', ''))[:14].encode('latin1', 'replace').decode('latin1'), border=1, align='L')
        pdf.cell(25, 7, str(row.get('Tam.', ''))[:10].encode('latin1', 'replace').decode('latin1'), border=1, align='C')
        pdf.cell(20, 7, str(abs(int(row.get('Qtd', 0)))), border=1, align='C')
        pdf.cell(20, 7, estado_texto[:10].encode('latin1', 'replace').decode('latin1'), border=1, align='C')
        pdf.ln()

    pdf.ln(10)
    pdf.set_font("Arial", '', 9)
    texto_declaracao = "Declaro que recebi o material descrito acima em perfeitas condições e me comprometo a utilizá-lo exclusivamente para as atividades relacionadas às minhas funções na NOVACAP."
    pdf.multi_cell(0, 5, texto_declaracao.encode('latin1', 'replace').decode('latin1'))
    pdf.ln(15)

    # Assinaturas
    y_assinatura = pdf.get_y()
    pdf.line(25, y_assinatura, 95, y_assinatura)
    pdf.line(115, y_assinatura, 185, y_assinatura)
    pdf.ln(3)
    pdf.set_font("Arial", '', 9)
    pdf.cell(95, 5, "Assinatura do Recebedor".encode('latin1', 'replace').decode('latin1'), align='C')
    pdf.cell(10, 5, "", align='C')
    pdf.cell(95, 5, "Assinatura do Almoxarife".encode('latin1', 'replace').decode('latin1'), align='C')
    pdf.ln(1)

    # Rodapé da folha A4
    pdf.set_y(-20)
    pdf.set_font("Arial", 'I', 8)
    texto_rodape = f"Sistema desenvolvido por Henrique Silva de Souza - {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    pdf.cell(0, 5, texto_rodape.encode('latin1', 'replace').decode('latin1'), ln=1, align='C')

    return pdf.output(dest='S')

def mostrar_pdf_inline(comprovante_pdf, altura=700):
    # Se o comprovante vier como string (caminho do arquivo), abre e lê os bytes
       # Se o comprovante vier como string (caminho ou texto), converte para bytes
    if isinstance(comprovante_pdf, str):
        if hasattr(comprovante_pdf, 'encode'):
            pdf_bytes = comprovante_pdf.encode('latin1')  # Codificação padrão de strings PDF
        else:
            with open(comprovante_pdf, "rb") as f:
                pdf_bytes = f.read()
    else:
        pdf_bytes = comprovante_pdf


    # Agora garante que passamos bytes puros para o base64
    base64_pdf = base64.b64encode(pdf_bytes).decode('utf-8')
    
    pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="{altura}" type="application/pdf"></iframe>'
    st.markdown(pdf_display, unsafe_allow_html=True)

# --- 5. CONTROLE DE ACESSO ---
if 'logado' not in st.session_state: 
    st.session_state.logado = False

if not st.session_state.logado:
    col1, col2, col3 = st.columns([1, 1.5, 1])
    with col2:
        exibir_logo(250)
        st.title("🔐 Acesso DSEL")
        u = st.text_input("Usuário").strip().upper()
        s = st.text_input("Senha", type="password")
        if st.button("Entrar"):
            with conectar() as conn:
                res = conn.execute("SELECT nome FROM usuarios WHERE UPPER(nome)=? AND senha=? AND status='Ativo'", (u, s)).fetchone()
            if res:
                st.session_state.logado, st.session_state.almoxarife = True, res[0]
                st.rerun()
            else: 
                st.error("Acesso negado.")
    rodape_henrique()
else:
    CORES = ["Cinza", "Amarelo", "Azul", "Branco", "Verde", "Preto", "Marrom", "Laranja", "Vermelho"]
    ITENS = ["Camiseta", "Calça", "Bota", "Luvas", "Chapéu", "Boné", "Cinto", "Meias"]
    TAMANHOS = ["Único","P", "M", "G", "GG", "XG","XGG"] + [str(i).zfill(2) for i in range(34, 57)]
    MODELOS = ["BRIM", "JEANS", "Padrão", "Botina", "Bontina Galocha", "luva de couro", "luva vaqueta", "Cano Alto", "Refletivo", "Elastano", "Tactel", "Camuflado"]

    with st.sidebar:
        exibir_logo(150)
        st.subheader(f"👤 {st.session_state.almoxarife}")
        opcoes = [
            "🏠 Painel Principal", 
            "📦 Saldo Geral", 
            "🔍 Consulta Individual", 
            "📥 Entrada (Fábrica)", 
            "📦 Entrega (Saída)", 
            "🔄 Retorno/Devolução", 
            "📑 Extrato por Item", 
            "📋 Extrato Geral", 
            "📊 Relatório Geral",
            "📅 Lançamentos Diários"
        ]
        if st.session_state.almoxarife == "ADMIN": 
            opcoes.append("⚙️ Gestão de Lançamentos")
            opcoes.append("🔧 Ajuste de Inventário")
            opcoes.append("🆕 Novo Cadastro")
            opcoes.append("✏️ Editar Usuários")
        menu = st.radio("Navegação", opcoes)
        if st.button("Sair"): 
            st.session_state.logado = False
            st.rerun()


    # --- 6. MÓDULOS ---
    if "Painel Principal" in menu:
        st.header("🏠 Detalhamento de Itens em Estoque")
        item_sel = st.selectbox("Selecione um item:", ["-- Selecione --"] + ITENS)
        if item_sel != "-- Selecione --":
            with conectar() as conn:
                df_det = pd.read_sql_query("SELECT modelo, cor, tamanho, estado, SUM(qtd) as 'Saldo' FROM estoque WHERE item=? GROUP BY modelo, cor, tamanho, estado HAVING SUM(qtd) != 0", conn, params=(item_sel,))
            if not df_det.empty:
                st.dataframe(df_det, use_container_width=True, hide_index=True)
                botoes_exportacao(df_det, f"saldo_{item_sel}", f"Saldo Detalhado - {item_sel}")
            else:
                st.info("Nenhum saldo encontrado para este item.")
        rodape_henrique()

    elif "Saldo Geral" in menu:
        st.header("📦 Saldo Geral Disponível")
        with conectar() as conn:
            df_saldo = pd.read_sql_query("SELECT item as 'Item', modelo as 'Modelo', cor as 'Cor', tamanho as 'Tam.', estado as 'Estado', SUM(qtd) as 'Saldo Atual' FROM estoque GROUP BY item, modelo, cor, tamanho, estado HAVING SUM(qtd) != 0 ORDER BY item ASC", conn)
        if not df_saldo.empty:
            st.dataframe(df_saldo, use_container_width=True, hide_index=True)
            st.info(f"💡 Total de peças físicas atualmente no estoque: **{df_saldo['Saldo Atual'].sum()}**")
            botoes_exportacao(df_saldo, "saldo_geral", "Saldo Geral de Estoque")
        else:
            st.info("Nenhum item com saldo positivo no estoque.")
        rodape_henrique()

    elif "Consulta Individual" in menu:
        st.header("🔍 Consulta por Reeducando")
        busca = st.text_input("Digite o Nome, CPF ou Matrícula:").strip().upper()
        if busca:
            with conectar() as conn:
                query = "SELECT data_hora as 'Data/Hora', item as 'Item', modelo as 'Modelo', cor as 'Cor', tamanho as 'Tam.', ABS(qtd) as 'Qtd. Retirada', estado as 'Estado', pessoa as 'Reeducando', almoxarife as 'Almoxarife' FROM estoque WHERE UPPER(pessoa) LIKE ? AND qtd < 0 ORDER BY id DESC"
                df_reeducando = pd.read_sql_query(query, conn, params=(f"%{busca}%",))
            if not df_reeducando.empty:
                st.dataframe(df_reeducando, use_container_width=True, hide_index=True)
                st.info(f"💡 Total de itens retirados: **{df_reeducando['Qtd. Retirada'].sum()}**")
                botoes_exportacao(df_reeducando, f"consulta_{busca}", "Consulta Individual de Reeducando")
            else:
                st.warning("Nenhum registro encontrado.")
        rodape_henrique()

    elif "Entrada (Fábrica)" in menu:
        st.header("📥 Entrada da Fábrica")
        with st.form("ent"):
            nf = st.text_input("Nota Fiscal")
            it, co, ta, mod = st.selectbox("Item", ITENS), st.selectbox("Cor", CORES), st.selectbox("Tamanho", TAMANHOS), st.selectbox("Modelo", MODELOS)
            qt = st.number_input("Quantidade", min_value=1)
            if st.form_submit_button("Salvar Entrada"):
                with conectar() as conn:
                    conn.execute("INSERT INTO estoque (item, cor, tamanho, modelo, qtd, estado, nf, pessoa, data_hora, almoxarife) VALUES (?,?,?,?,?,?,?,?,?,?)", 
                                 (it, co, ta, mod, qt, "Novo", nf, "Fábrica", datetime.now().strftime("%d/%m/%Y %H:%M"), st.session_state.almoxarife))
                st.success("Registrado!")
                st.rerun()
        rodape_henrique()

    elif "Entrega (Saída)" in menu:
        st.header("📤 Entrega de Fardamento")
        col1, col2, col3 = st.columns(3)
        with col1: it = st.selectbox("Item", ITENS, index=ITENS.index("Calça") if "Calça" in ITENS else 0)
        with col2: co = st.selectbox("Cor", CORES)
        with col3: ta = st.selectbox("Tamanho", TAMANHOS)
        col4, col5 = st.columns(2)
        with col4: mod = st.selectbox("Modelo", MODELOS)
        with col5: est = st.radio("Origem:", ["Novo", "Usado/Reuso"], horizontal=True)

        saldo = consultar_saldo(it, co, ta, mod, est)
        if saldo > 0:
            st.success(f"✅ Saldo disponível: {saldo}")
        else:
            st.error(f"❌ Estoque ZERADO para {it} {mod}")

        with st.form("confirmar_saida"):
            dest = st.text_input("Recebedor").upper()
            qt_saida = st.number_input("Quantidade", min_value=1, max_value=max(1, int(saldo)))
            if st.form_submit_button("Finalizar Entrega"):
                saldo_atualizado = consultar_saldo(it, co, ta, mod, est)
                if saldo_atualizado <= 0 or qt_saida > saldo_atualizado:
                    st.error("Não há saldo suficiente disponível no momento do envio!")
                elif not dest:
                    st.warning("Preencha o recebedor.")
                else:
                    with conectar() as conn:
                        conn.execute("INSERT INTO estoque (item, cor, tamanho, modelo, qtd, estado, pessoa, data_hora, almoxarife) VALUES (?,?,?,?,?,?,?,?,?)", 
                                     (it, co, ta, mod, -int(qt_saida), est, dest, datetime.now().strftime("%d/%m/%Y %H:%M"), st.session_state.almoxarife))
                    st.success("Saída registrada com sucesso!")
                    st.rerun()
        rodape_henrique()

    elif "Retorno/Devolução" in menu:
        st.header("🔄 Retorno / Devolução de Item")
        with st.form("devolucao"):
            it, co, ta, mod = st.selectbox("Item", ITENS), st.selectbox("Cor", CORES), st.selectbox("Tamanho", TAMANHOS), st.selectbox("Modelo", MODELOS)
            est = st.radio("Estado de retorno:", ["Usado/Reuso", "Novo"], horizontal=True)
            qt_ret = st.number_input("Quantidade Retornando", min_value=1)
            origem = st.text_input("Quem está devolvendo").upper()
            if st.form_submit_button("Registrar Retorno"):
                if not origem:
                    st.warning("Informe o reeducando.")
                else:
                    with conectar() as conn:
                        conn.execute("INSERT INTO estoque (item, cor, tamanho, modelo, qtd, estado, pessoa, data_hora, almoxarife) VALUES (?,?,?,?,?,?,?,?,?)", 
                                     (it, co, ta, mod, int(qt_ret), est, f"Devolução: {origem}", datetime.now().strftime("%d/%m/%Y %H:%M"), st.session_state.almoxarife))
                    st.success("Retorno registrado!")
                    st.rerun()
        rodape_henrique()

    elif "Extrato por Item" in menu:
        st.header("📑 Extrato de Movimentação por Item")
        col1, col2, col3 = st.columns(3)
        with col1: it = st.selectbox("Item", ITENS)
        with col2: co = st.selectbox("Cor", CORES)
        with col3: ta = st.selectbox("Tamanho", TAMANHOS)
        col4, col5 = st.columns(2)
        with col4: mod = st.selectbox("Modelo", MODELOS)
        with col5: est = st.radio("Estado:", ["Novo", "Usado/Reuso"], horizontal=True)

        with conectar() as conn:
            df_extrato = pd.read_sql_query("SELECT data_hora as 'Data/Hora', nf as 'NF', pessoa as 'Origem/Destino', qtd as 'Qtd' FROM estoque WHERE item=? AND cor=? AND tamanho=? AND modelo=? AND estado=? ORDER BY id ASC", conn, params=(it, co, ta, mod, est))

        if not df_extrato.empty:
            df_extrato['Saldo Acumulado'] = df_extrato['Qtd'].cumsum()
            st.dataframe(df_extrato, use_container_width=True, hide_index=True)
            st.metric(label="Saldo Final Disponível", value=f"{df_extrato['Saldo Acumulado'].iloc[-1]} unidades")
            botoes_exportacao(df_extrato, f"extrato_{it}_{mod}", f"Extrato por Item - {it} {mod}")
        else:
            st.info("Nenhuma movimentação encontrada.")
        rodape_henrique()

    elif "Extrato Geral" in menu:
        st.header("📋 Extrato Geral de Movimentações")
        st.write("Histórico consolidado de movimentações para visualização e impressão.")
        with conectar() as conn:
            df_extrato_geral = pd.read_sql_query("SELECT data_hora as 'Data/Hora', item as 'Item', modelo as 'Modelo', cor as 'Cor', tamanho as 'Tam.', qtd as 'Qtd', estado as 'Estado', nf as 'NF', pessoa as 'Origem/Destino', almoxarife as 'Almoxarife' FROM estoque ORDER BY id DESC", conn)

        if not df_extrato_geral.empty:
            botoes_exportacao(df_extrato_geral, "extrato_geral", "Extrato Geral de Movimentações")
            st.write("---")
            ver_preview = st.checkbox("👁️ Visualizar como ficará impresso")
            if ver_preview:
                st.subheader("📄 Pré-visualização da Impressão")
                st.markdown("<div style='background-color: #ffffff; color: #111111; padding: 25px; border-radius: 5px; border: 2px solid #333; font-family: Courier New, monospace;'><h3 style='text-align: center; margin-bottom: 0;'>NOVACAP - DIVISÃO DE SEGURANÇA, LIMPEZA E CONSERVAÇÃO (DSEL)</h3><p style='text-align: center; margin-top: 5px;'><b>Relatório consolidado de Extrato Geral</b></p><p style='text-align: right; font-size: 12px;'>Data de emissão: " + datetime.now().strftime('%d/%m/%Y %H:%M') + "</p><hr style='border-top: 1px dashed black; margin-bottom: 15px;'></div>", unsafe_allow_html=True)
                st.dataframe(df_extrato_geral, use_container_width=True, hide_index=True)
            st.write("### 📋 Tabela de Dados no Sistema")
            st.dataframe(df_extrato_geral, use_container_width=True, hide_index=True)
        else:
            st.info("Nenhuma movimentação registrada no sistema.")
        rodape_henrique()

    elif "Relatório Geral" in menu:
        st.header("📊 Movimentação Completa (Tabela do Banco)")
        with conectar() as conn:
            df_geral = pd.read_sql_query("SELECT * FROM estoque ORDER BY id DESC", conn)
        if not df_geral.empty:
            st.dataframe(df_geral, use_container_width=True, hide_index=True)
            botoes_exportacao(df_geral, "relatorio_geral", "Relatório Geral Completo")
        else:
            st.info("Nenhuma movimentação registrada.")
        rodape_henrique()

    elif "Lançamentos Diários" in menu:
        st.header("📅 Lançamentos do Dia")
    
    # 1. Captura a data do calendário
        data_calendario = st.date_input("Selecione a data para filtrar:", date.today(), format="DD/MM/YYYY")
    
    # 2. Formata a data em texto
        data_sel_str = data_calendario.strftime("%d/%m/%Y")
    
    # 3. EXIBE A DATA
        st.subheader(f"Data: {data_sel_str}")
    
        with conectar() as conn:
            df_diario = pd.read_sql_query(
                "SELECT data_hora as 'Data/Hora', item as 'Item', modelo as 'Modelo', cor as 'Cor', tamanho as 'Tam.', qtd as 'Qtd', estado as 'Estado', nf as 'NF', pessoa as 'Origem/Destino', almoxarife as 'Almoxarife' FROM estoque WHERE data_hora LIKE ? ORDER BY id DESC", 
                conn, 
                params=(f"{data_sel_str}%",)
            )

        if not df_diario.empty:
            st.success(f"✅ {len(df_diario)} lançamento(s) encontrado(s) para {data_sel_str}.")
            st.dataframe(df_diario, use_container_width=True, hide_index=True)

            entradas = df_diario[df_diario['Qtd'] > 0]['Qtd'].sum()
            saidas = abs(df_diario[df_diario['Qtd'] < 0]['Qtd'].sum())
            col1, col2 = st.columns(2)
            col1.metric("📥 Total de Entradas", f"{entradas} un")
            col2.metric("📤 Total de Saídas", f"{saidas} un")

            botoes_exportacao(df_diario, f"lancamentos_diarios_{data_sel_str.replace('/','')}", "Lançamentos Diários - DSEL")

            st.divider()
            st.subheader("🖨️ Comprovante de Entrega por Pessoa")
            st.write(f"Selecione uma pessoa para gerar um comprovante de entrega individualizado (PDF) com os itens que ela recebeu em {data_sel_str}.")

            saidas_hoje = df_diario[df_diario['Qtd'] < 0]
            saidas_hoje = saidas_hoje[
                ~saidas_hoje['Origem/Destino'].astype(str).str.upper().isin(['FÁBRICA', 'AJUSTE']) &
                ~saidas_hoje['Origem/Destino'].astype(str).str.upper().str.startswith('DEVOLUÇÃO')
            ]

            if not saidas_hoje.empty:
                pessoas_unicas = sorted(saidas_hoje['Origem/Destino'].dropna().unique())
                pessoa_sel = st.selectbox("Selecione o recebedor:", ["-- Selecione --"] + list(pessoas_unicas))

                if pessoa_sel != "-- Selecione --":
                    df_pessoa = saidas_hoje[saidas_hoje['Origem/Destino'] == pessoa_sel][
                        ['Item', 'Modelo', 'Cor', 'Tam.', 'Qtd', 'Estado']
                    ].copy()
                    df_pessoa['Qtd'] = df_pessoa['Qtd'].abs()

                    st.write(f"**Itens entregues em {data_sel_str} para {pessoa_sel}:**")
                    st.dataframe(df_pessoa, use_container_width=True, hide_index=True)

                    comprovante_pdf = gerar_comprovante_entrega(pessoa_sel, df_pessoa, st.session_state.almoxarife)

                    st.write("---")
                    st.subheader("👁️ Pré-visualização do Comprovante")
                    mostrar_pdf_inline(comprovante_pdf, altura=700)

                    st.download_button(
                        label="📄 Baixar Comprovante de Entrega (PDF)",
                        data=comprovante_pdf,
                        file_name=f"comprovante_entrega_{pessoa_sel.replace(' ','_')}_{data_sel_str.replace('/','')}.pdf",
                        mime="application/pdf"
                    )

                    st.info("💡 O PDF gerado inclui: logo NOVACAP, cabeçalho DSEL, lista de itens, data, nome do almoxarife, declaração de recebimento, campos para assinatura e rodapé com crédito do desenvolvedor.")
            else:
                st.info(f"Nenhuma entrega (saída) registrada em {data_sel_str} para geração de comprovantes.")
        else:
            st.info(f"Nenhum lançamento registrado em {data_sel_str}.")
        rodape_henrique()


    elif "Gestão de Lançamentos" in menu and st.session_state.almoxarife == "ADMIN":
        st.header("⚙️ Gestão de Lançamentos (Modo Supervisor)")
        st.warning("⚠️ Atenção: A exclusão ou edição de registros altera o saldo do estoque permanentemente.")
        with conectar() as conn:
            df_admin = pd.read_sql_query("SELECT * FROM estoque ORDER BY id DESC", conn)

        if not df_admin.empty:
            st.subheader("🔍 1. Encontrar Lançamento")
            termo_busca = st.text_input("Digite o nome do Recebedor/Reeducando, ID, NF ou Item para filtrar:").strip().upper()

            if termo_busca:
                df_filtrado = df_admin[
                    df_admin['pessoa'].astype(str).str.upper().str.contains(termo_busca, na=False) |
                    df_admin['item'].astype(str).str.upper().str.contains(termo_busca, na=False) |
                    df_admin['nf'].astype(str).str.upper().str.contains(termo_busca, na=False) |
                    df_admin['id'].astype(str).str.contains(termo_busca, na=False)
                ]
            else:
                df_filtrado = df_admin

            if not df_filtrado.empty:
                lista_opcoes = {
                    f"ID: {row['id']} | Data: {row['data_hora']} | {row['item']} ({row['modelo']}) | Qtd: {row['qtd']} | Destino: {row['pessoa']} | Estado: {row['estado']}": row['id']
                    for idx, row in df_filtrado.iterrows()
                }
                selecionado_rotulo = st.selectbox("Selecione o registro exato para alterar:", ["-- Selecione --"] + list(lista_opcoes.keys()))

                if selecionado_rotulo != "-- Selecione --":
                    id_registro_banco = lista_opcoes[selecionado_rotulo]
                    reg_atual = df_admin[df_admin['id'] == id_registro_banco].iloc[0]
                    st.divider()
                    st.subheader(f"📝 Formulário de Alteração - Registro Interno")

                    col_form1, col_form2, col_form3 = st.columns(3)
                    with col_form1:
                        novo_item = st.selectbox("Corrigir Item:", ITENS, index=ITENS.index(reg_atual['item']) if reg_atual['item'] in ITENS else 0)
                        nova_cor = st.selectbox("Corrigir Cor:", CORES, index=CORES.index(reg_atual['cor']) if reg_atual['cor'] in CORES else 0)
                    with col_form2:
                        novo_modelo = st.selectbox("Corrigir Modelo:", MODELOS, index=MODELOS.index(reg_atual['modelo']) if reg_atual['modelo'] in MODELOS else 0)
                        novo_tamanho = st.selectbox("Corrigir Tam:", TAMANHOS, index=TAMANHOS.index(reg_atual['tamanho']) if reg_atual['tamanho'] in TAMANHOS else 0)
                    with col_form3:
                        nova_qtd = st.number_input("Corrigir Qtd (Valores negativos indicam Saídas):", value=int(reg_atual['qtd']))
                        novo_destino = st.text_input("Corrigir Responsável/Destino:", value=str(reg_atual['pessoa'])).upper()

                    st.write("---")
                    st.subheader("🔄 Alterar Estado do Item")
                    estado_atual = reg_atual['estado'] if reg_atual['estado'] else "Novo"
                    novo_estado = st.radio(
                        "Selecione o novo estado:", 
                        ["Novo", "Usado/Reuso"], 
                        horizontal=True,
                        index=0 if estado_atual == "Novo" else 1
                    )
                    if estado_atual != novo_estado:
                        st.info(f"O estado será alterado de **{estado_atual}** para **{novo_estado}**")

                    st.write("---")
                    col_btn1, col_btn2, col_btn3 = st.columns(3)

                    with col_btn1:
                        if st.button("💾 Salvar Alterações", type="primary"):
                            with conectar() as conn:
                                conn.execute("UPDATE estoque SET item=?, modelo=?, cor=?, tamanho=?, qtd=?, pessoa=?, estado=? WHERE id=?", 
                                             (novo_item, novo_modelo, nova_cor, novo_tamanho, int(nova_qtd), novo_destino, novo_estado, int(id_registro_banco)))
                                conn.commit()
                            st.success("Lançamento atualizado com sucesso!")
                            st.rerun()

                    with col_btn2:
                        if st.button("🚨 Confirmar Exclusão Definitiva", type="secondary"):
                            with conectar() as conn:
                                conn.execute("DELETE FROM estoque WHERE id = ?", (int(id_registro_banco),))
                                conn.commit()
                            st.warning("Lançamento apagado permanentemente.")
                            st.rerun()

                    with col_btn3:
                        if st.button("❌ Cancelar Operação"):
                            st.info("Ação cancelada.")
                            st.rerun()
            else:
                st.error("Nenhum lançamento corresponde aos termos pesquisados.")
            st.write("### 📋 Histórico Completo de Auditoria")
            st.dataframe(df_admin, use_container_width=True, hide_index=True)
            botoes_exportacao(df_admin, "gestao_lancamentos", "Gestão de Lançamentos - Auditoria")
        else:
            st.info("Nenhuma movimentação registrada no sistema.")
        rodape_henrique()

    elif "Ajuste de Inventário" in menu and st.session_state.almoxarife == "ADMIN":
        st.header("🔧 Ajuste de Balanço de Inventário")
        st.write("Utilize esta aba para forçar o saldo do sistema a bater exatamente com o estoque físico contado.")

        col_adj1, col_adj2, col_adj3 = st.columns(3)
        with col_adj1: item_adj = st.selectbox("Item para Ajustar", ITENS)
        with col_adj2: cor_adj = st.selectbox("Cor para Ajustar", CORES)
        with col_adj3: tam_adj = st.selectbox("Tam para Ajustar", TAMANHOS)

        col_adj4, col_adj5 = st.columns(2)
        with col_adj4: mod_adj = st.selectbox("Modelo para Ajustar", MODELOS)
        with col_adj5: est_adj = st.radio("Estado do Item:", ["Novo", "Usado/Reuso"], horizontal=True)

        saldo_sistema = consultar_saldo(item_adj, cor_adj, tam_adj, mod_adj, est_adj)
        st.metric(label="Saldo Atual no Sistema", value=f"{saldo_sistema} unidades")

        with st.form("form_ajuste"):
            qtd_real_fisica = st.number_input("Quantidade Real Encontrada no Armário (Físico):", min_value=0, value=int(saldo_sistema))
            motivo_ajuste = st.text_input("Motivo do Ajuste (Ex: Inventário Geral 2026, Correção de quebra)").upper()

            if st.form_submit_button("🔨 Aplicar Ajuste de Estoque"):
                diferenca = qtd_real_fisica - saldo_sistema
                if diferenca == 0:
                    st.info("O valor digitado é igual ao do sistema. Nenhum ajuste necessário.")
                elif not motivo_ajuste:
                    st.warning("Descreva o motivo do ajuste para fins de auditoria.")
                else:
                    with conectar() as conn:
                        conn.execute("INSERT INTO estoque (item, cor, tamanho, modelo, qtd, estado, nf, pessoa, data_hora, almoxarife) VALUES (?,?,?,?,?,?,?,?,?,?)", 
                                     (item_adj, cor_adj, tam_adj, mod_adj, int(diferenca), est_adj, "AJUSTE", f"INV: {motivo_ajuste}", 
                                      datetime.now().strftime("%d/%m/%Y %H:%M"), st.session_state.almoxarife))
                        conn.commit()
                    st.success(f"Estoque ajustado! Lançamento de balanceamento de {diferenca} unidades registrado.")
                    st.rerun()
        rodape_henrique()

    elif "Novo Cadastro" in menu and st.session_state.almoxarife == "ADMIN":
        st.header("🆕 Novo Cadastro de Usuários do Sistema")
        with st.form("form_novo_usuario"):
            novo_nome = st.text_input("Nome do Usuário/Almoxarife").strip().upper()
            nova_senha = st.text_input("Senha de Acesso", type="password")
            novo_status = st.selectbox("Status Inicial", ["Ativo", "Inativo"])

            if st.form_submit_button("Salvar Usuário"):
                if not novo_nome or not nova_senha:
                    st.warning("Preencha todos os campos obrigatórios.")
                else:
                    try:
                        with conectar() as conn:
                            conn.execute("INSERT INTO usuarios (nome, senha, status) VALUES (?, ?, ?)", (novo_nome, nova_senha, novo_status))
                            conn.commit()
                        st.success(f"Usuário {novo_nome} cadastrado com sucesso!")
                    except sqlite3.IntegrityError:
                        st.error("Erro: Já existe um usuário cadastrado com esse nome.")
        rodape_henrique()

    elif "Editar Usuários" in menu and st.session_state.almoxarife == "ADMIN":
        st.header("✏️ Gerenciamento e Edição de Usuários")
        with conectar() as conn:
            df_usuarios = pd.read_sql_query("SELECT id, nome, status FROM usuarios ORDER BY nome ASC", conn)

        if not df_usuarios.empty:
            st.subheader("Lista de Usuários Cadastrados")
            st.dataframe(df_usuarios, use_container_width=True, hide_index=True)
            botoes_exportacao(df_usuarios, "usuarios_cadastrados", "Usuários do Sistema")

            st.write("---")
            st.subheader("Alterar Dados de um Usuário")
            usuario_selecionado = st.selectbox("Selecione o usuário para modificar:", df_usuarios["nome"].tolist())

            with conectar() as conn:
                dados_user = conn.execute("SELECT senha, status FROM usuarios WHERE nome = ?", (usuario_selecionado,)).fetchone()

            if dados_user:
                with st.form("form_edicao_usuario"):
                    ed_senha = st.text_input("Nova Senha", value=str(dados_user[0]))
                    ed_status = st.selectbox("Alterar Status", ["Ativo", "Inativo"], index=0 if dados_user[1] == "Ativo" else 1)

                    if st.form_submit_button("Confirmar Alterações"):
                        with conectar() as conn:
                            conn.execute("UPDATE usuarios SET senha = ?, status = ? WHERE nome = ?", (ed_senha, ed_status, usuario_selecionado))
                            conn.commit()
                        st.success(f"Dados do usuário {usuario_selecionado} atualizados!")
                        st.rerun()
        else:
            st.info("Nenhum usuário cadastrado no sistema.")
        rodape_henrique()
